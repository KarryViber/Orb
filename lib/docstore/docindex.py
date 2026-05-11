#!/usr/bin/env python3
from __future__ import annotations

import argparse
import hashlib
import json
import logging
import re
import sqlite3
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

try:
    import docx  # type: ignore
except Exception:  # pragma: no cover
    docx = None

try:
    import pdfplumber  # type: ignore
except Exception:  # pragma: no cover
    pdfplumber = None

logging.getLogger("pdfminer").setLevel(logging.ERROR)
logging.getLogger("pdfplumber").setLevel(logging.ERROR)

import os as _os
_os.environ.setdefault("KMP_DUPLICATE_LIB_OK", "TRUE")
_os.environ.setdefault("TOKENIZERS_PARALLELISM", "false")

# Paths MUST be configured via env vars — no hardcoded fallbacks.
_projects_root = _os.environ.get("DOC_PROJECTS_ROOT", "")
if not _projects_root:
    PROJECTS_ROOT = None
    REGISTRY_PATH = None
else:
    PROJECTS_ROOT = Path(_projects_root)
    REGISTRY_PATH = PROJECTS_ROOT / "registry.md"

# DOC_ROOT enables the wide-scan mode (work/). When unset, falls back to
# legacy behavior scanning only PROJECTS_ROOT/*/{00_source,...}.
_doc_root = _os.environ.get("DOC_ROOT", "")
DOC_ROOT = Path(_doc_root) if _doc_root else None

_references_root = _os.environ.get("DOC_REFERENCES_ROOT", "")
REFERENCES_ROOT = Path(_references_root) if _references_root else None

_doc_index_db = _os.environ.get("DOC_INDEX_DB", "")
DEFAULT_DB_PATH = Path(_doc_index_db) if _doc_index_db else None

_docstore_wide_map_json = _os.environ.get("DOCSTORE_WIDE_MAP_JSON", "")
DOCSTORE_WIDE_MAP_JSON = Path(_docstore_wide_map_json) if _docstore_wide_map_json else None
MAX_CHUNK_CHARS = 1200
MIN_CHUNK_CHARS = 300
OVERLAP_CHARS = 180

# Legacy mode: per-project fixed subdirs with per-subdir allowed suffixes.
INCLUDE_DIR_EXTS = {
    "00_source": {".md", ".docx", ".pdf"},
    "01_meetings": {".md", ".docx", ".pdf"},
    "02_draft": {".md", ".docx", ".pdf"},
    "03_delivery": {".md", ".docx", ".pdf"},
}

# Wide-scan mode: uniform allowed suffixes across the whole tree.
ALLOWED_SUFFIXES = {".md", ".docx", ".pdf"}

# Directory segment names that exclude the entire subtree in wide-scan mode.
# Draft-family dirs are excluded to reduce noise in agent retrieval.
EXCLUDE_DIR_SEGMENTS = {
    "02_draft",
    "01_briefs",
    "02_drafts",
    "03_publish_packets",
    "receipts",
    "slides",
    "node_modules",
    ".git",
    ".obsidian",
    "brand-kit",
    "__pycache__",
}

# Known doc_type segments inside a project directory.
KNOWN_PROJECT_DOC_TYPES = {"00_source", "01_meetings", "03_delivery"}

EXCLUDE_SUFFIXES = {
    ".drawio",
    ".bkp",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".gif",
    ".svg",
    ".pptx",
    ".xlsx",
    ".csv",
    ".js",
    ".py",
}

EXCLUDE_NAME_SET = {
    ".DS_Store",
    "CLAUDE.md",          # agent 指示文件，非业务资料
    "INDEX.md",            # 目录索引，非业务资料
    "INDEX_TEMPLATE.md",   # 模板占位
    "registry.md",         # slug 映射表
}
EXCLUDE_PATH_PARTS = {"slides", "node_modules", ".git", ".obsidian", "brand-kit", ".pytest_cache"}
EXCLUDE_PATH_SUBSTRINGS = (
    "/archived/",
    "_archived_",
    "/voice-booking-demo/code/",  # 代码仓内的 README/prompts 非业务文档
)

# doc_type ranking boost: multiplied with bm25 (negative scores).
# Higher multiplier = more negative = ranked higher.
DOC_TYPE_BOOST = {
    "03_delivery": 1.5,   # 最高优先：确定性交付物
    "published": 1.3,     # articles/04_published / 05_archive
    "00_source": 1.2,     # 高优先：客户原件
    "01_meetings": 1.0,   # 正常：会议纪要
    "reference": 1.0,     # partners/psr/杂项正式材料
    "02_draft": 0.6,      # 降权：编辑中草稿
}

EMBEDDING_DIM = 1024
EMBEDDING_MODEL_NAME = "BAAI/bge-m3"
EMBEDDING_BATCH_SIZE = 4

SCHEMA_SQL = """
PRAGMA journal_mode=WAL;
PRAGMA foreign_keys=ON;

CREATE TABLE IF NOT EXISTS documents (
    id INTEGER PRIMARY KEY,
    path TEXT NOT NULL UNIQUE,
    slug TEXT NOT NULL,
    project_dir TEXT NOT NULL,
    doc_type TEXT NOT NULL,
    title TEXT NOT NULL,
    ext TEXT NOT NULL,
    mtime_ns INTEGER NOT NULL,
    size_bytes INTEGER NOT NULL,
    content_sha256 TEXT NOT NULL,
    indexed_at TEXT NOT NULL
);

CREATE TABLE IF NOT EXISTS chunks (
    id INTEGER PRIMARY KEY,
    doc_id INTEGER NOT NULL REFERENCES documents(id) ON DELETE CASCADE,
    chunk_index INTEGER NOT NULL,
    path TEXT NOT NULL,
    slug TEXT NOT NULL,
    doc_type TEXT NOT NULL,
    title TEXT NOT NULL,
    section TEXT NOT NULL,
    content TEXT NOT NULL,
    embedding BLOB,
    priority_layer TEXT NOT NULL DEFAULT 'L5',
    UNIQUE(doc_id, chunk_index)
);

CREATE VIRTUAL TABLE IF NOT EXISTS chunks_fts USING fts5(
    path,
    slug,
    doc_type,
    title,
    section,
    content,
    tokenize = 'trigram'
);

CREATE TABLE IF NOT EXISTS index_meta (
    key TEXT PRIMARY KEY,
    value TEXT NOT NULL
);

CREATE INDEX IF NOT EXISTS idx_documents_slug_doc_type ON documents(slug, doc_type);
CREATE INDEX IF NOT EXISTS idx_chunks_doc_id ON chunks(doc_id);
CREATE INDEX IF NOT EXISTS idx_chunks_slug_doc_type ON chunks(slug, doc_type);
"""


@dataclass(frozen=True)
class CandidateFile:
    path: Path
    slug: str
    project_dir: str
    doc_type: str


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    title_hint: str


def now_iso() -> str:
    return datetime.now().astimezone().isoformat(timespec="seconds")


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def normalize_text(text: str) -> str:
    text = text.replace("\u0000", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def load_registry_dir_slug_map() -> dict[str, str]:
    if not REGISTRY_PATH or not REGISTRY_PATH.exists():
        return {}

    current_section = "projects"
    mapping: dict[str, str] = {}
    for raw_line in REGISTRY_PATH.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if line.startswith("## "):
            lower = line.lower()
            if "internal registry" in lower:
                current_section = "internal"
            elif "partner registry" in lower:
                current_section = "partner"
            else:
                current_section = "projects"
            continue
        if current_section != "projects":
            continue
        if not line.startswith("|"):
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if len(cells) < 3:
            continue
        if cells[0] in {"slug", "------"}:
            continue
        slug, directory = cells[0], cells[1]
        if slug and directory:
            mapping[directory] = slug
    return mapping


def should_index(path: Path, doc_type: str) -> bool:
    if path.name in EXCLUDE_NAME_SET:
        return False
    if any(part in EXCLUDE_PATH_PARTS for part in path.parts):
        return False
    posix = path.as_posix()
    if any(fragment in posix for fragment in EXCLUDE_PATH_SUBSTRINGS):
        return False
    suffix = path.suffix.lower()
    if suffix in EXCLUDE_SUFFIXES:
        return False
    allowed = INCLUDE_DIR_EXTS.get(doc_type, set())
    return suffix in allowed


def load_docstore_wide_map_rules() -> list[dict[str, Any]]:
    """Load wide-mode path mapping rules from DOCSTORE_WIDE_MAP_JSON.

    Wide mode intentionally has no built-in organization-specific paths. Set
    DOCSTORE_WIDE_MAP_JSON to a JSON file with a top-level `rules` array, or to
    a JSON array directly. See docs/docstore-wide-map.example.json.
    """
    if not DOCSTORE_WIDE_MAP_JSON:
        return []
    if not DOCSTORE_WIDE_MAP_JSON.exists():
        raise FileNotFoundError(f"DOCSTORE_WIDE_MAP_JSON not found: {DOCSTORE_WIDE_MAP_JSON}")
    payload = json.loads(DOCSTORE_WIDE_MAP_JSON.read_text(encoding="utf-8"))
    rules = payload.get("rules") if isinstance(payload, dict) else payload
    if not isinstance(rules, list):
        raise ValueError("DOCSTORE_WIDE_MAP_JSON must contain a rules array")
    out: list[dict[str, Any]] = []
    for index, rule in enumerate(rules):
        if not isinstance(rule, dict):
            raise ValueError(f"wide map rule {index} must be an object")
        match = rule.get("match")
        if not isinstance(match, list) or not all(isinstance(item, str) for item in match):
            raise ValueError(f"wide map rule {index} must include a string match array")
        out.append(rule)
    return out


def _match_wide_rule(pattern: list[str], rel_parts: tuple[str, ...]) -> list[str] | None:
    captures: list[str] = []
    part_index = 0
    for token in pattern:
        if token == "**":
            captures.extend(rel_parts[part_index:])
            part_index = len(rel_parts)
            break
        if part_index >= len(rel_parts):
            return None
        current = rel_parts[part_index]
        if token == "*":
            captures.append(current)
        elif token != current:
            return None
        part_index += 1
    if part_index != len(rel_parts):
        return None
    return captures


def _render_wide_template(template: object, captures: list[str], dir_to_slug: dict[str, str]) -> str:
    if not isinstance(template, str):
        raise ValueError("wide map output templates must be strings")

    def replace(match: re.Match[str]) -> str:
        expression = match.group(1)
        parts = expression.split("|")
        try:
            value = captures[int(parts[0])]
        except (ValueError, IndexError) as err:
            raise ValueError(f"invalid wide map capture reference: {expression}") from err
        for op in parts[1:]:
            if op == "registry":
                value = dir_to_slug.get(value, value)
            elif op.startswith("known_project_doc_type:"):
                fallback = op.split(":", 1)[1]
                value = value if value in KNOWN_PROJECT_DOC_TYPES else fallback
            else:
                raise ValueError(f"unsupported wide map template operator: {op}")
        return value

    return re.sub(r"\{([^{}]+)\}", replace, template)


def derive_ids_wide(
    rel_parts: tuple[str, ...],
    dir_to_slug: dict[str, str],
    mapping_rules: list[dict[str, Any]] | None = None,
) -> tuple[str, str, str] | None:
    """Map a path relative to DOC_ROOT to (slug, project_dir, doc_type).

    Returns None if no configured wide-map rule matches the path.
    """
    if not rel_parts:
        return None
    rules = load_docstore_wide_map_rules() if mapping_rules is None else mapping_rules
    for rule in rules:
        captures = _match_wide_rule(rule["match"], rel_parts)
        if captures is None:
            continue
        slug = _render_wide_template(rule.get("slug"), captures, dir_to_slug)
        project_dir = _render_wide_template(rule.get("project_dir"), captures, dir_to_slug)
        doc_type_template = rule.get("doc_type", "reference")
        doc_type = _render_wide_template(doc_type_template, captures, dir_to_slug)
        return slug, project_dir, doc_type
    return None


def should_index_wide(path: Path) -> bool:
    if path.name in EXCLUDE_NAME_SET:
        return False
    if path.name.startswith("~$"):  # Office lock files
        return False
    parts = path.parts
    if any(part in EXCLUDE_DIR_SEGMENTS for part in parts):
        return False
    if any(part in EXCLUDE_PATH_PARTS for part in parts):
        return False
    posix = path.as_posix()
    if any(fragment in posix for fragment in EXCLUDE_PATH_SUBSTRINGS):
        return False
    suffix = path.suffix.lower()
    if suffix in EXCLUDE_SUFFIXES:
        return False
    return suffix in ALLOWED_SUFFIXES


def iter_candidate_files(slug_filter: str | None = None) -> Iterable[CandidateFile]:
    dir_to_slug = load_registry_dir_slug_map()

    if DOC_ROOT and DOC_ROOT.exists():
        for path in sorted(DOC_ROOT.rglob("*")):
            if not path.is_file():
                continue
            if not should_index_wide(path):
                continue
            rel = path.relative_to(DOC_ROOT).parts
            ids = derive_ids_wide(rel, dir_to_slug)
            if not ids:
                continue
            slug, project_dir, doc_type = ids
            if slug_filter and slug != slug_filter:
                continue
            yield CandidateFile(path=path, slug=slug, project_dir=project_dir, doc_type=doc_type)
    else:
        # Legacy mode: PROJECTS_ROOT/*/docType/**
        if PROJECTS_ROOT and PROJECTS_ROOT.exists():
            for project_dir in sorted(PROJECTS_ROOT.iterdir()):
                if not project_dir.is_dir() or project_dir.name.startswith("."):
                    continue
                slug = dir_to_slug.get(project_dir.name, project_dir.name)
                if slug_filter and slug != slug_filter:
                    continue
                for doc_type in INCLUDE_DIR_EXTS:
                    base = project_dir / doc_type
                    if not base.exists():
                        continue
                    for path in sorted(base.rglob("*")):
                        if path.is_file() and should_index(path, doc_type):
                            yield CandidateFile(path=path, slug=slug, project_dir=project_dir.name, doc_type=doc_type)

    if REFERENCES_ROOT and REFERENCES_ROOT.exists():
        path = REFERENCES_ROOT / "bookmarks.md"
        if path.is_file() and (not slug_filter or slug_filter == "references"):
            yield CandidateFile(path=path, slug="references", project_dir="references", doc_type="bookmark")


def extract_title(text: str, fallback: str) -> str:
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            title = stripped.lstrip("#").strip()
            if title:
                return title
    return fallback


def normalize_section(label: str | None) -> str:
    if not label:
        return "(root)"
    return label.strip() or "(root)"


def chunk_text(text: str) -> list[tuple[str, str]]:
    sections: list[tuple[str, str]] = []
    current_heading = "(root)"
    current_lines: list[str] = []

    def flush_section() -> None:
        body = "\n".join(current_lines).strip()
        if body:
            sections.append((normalize_section(current_heading), body))
        current_lines.clear()

    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            flush_section()
            current_heading = stripped.lstrip("#").strip() or "(root)"
            continue
        current_lines.append(line)
    flush_section()

    if not sections:
        collapsed = text.strip()
        if collapsed:
            sections = [("(root)", collapsed)]

    chunks: list[tuple[str, str]] = []
    for section, body in sections:
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n+", body) if p.strip()]
        if not paragraphs:
            continue
        buffer = ""
        for paragraph in paragraphs:
            candidate = paragraph if not buffer else f"{buffer}\n\n{paragraph}"
            if len(candidate) <= MAX_CHUNK_CHARS:
                buffer = candidate
                continue
            if buffer:
                chunks.append((section, buffer))
                if len(buffer) >= MIN_CHUNK_CHARS:
                    overlap = buffer[-OVERLAP_CHARS:]
                    overlap = overlap.split("\n", 1)[-1].strip() if "\n" in overlap else overlap.strip()
                    buffer = overlap if overlap else ""
                else:
                    buffer = ""
            while len(paragraph) > MAX_CHUNK_CHARS:
                slice_text = paragraph[:MAX_CHUNK_CHARS].strip()
                chunks.append((section, slice_text))
                paragraph = paragraph[MAX_CHUNK_CHARS - OVERLAP_CHARS :].strip()
            buffer = paragraph if not buffer else f"{buffer}\n\n{paragraph}"
        if buffer.strip():
            chunks.append((section, buffer.strip()))

    return [(section, content) for section, content in chunks if content.strip()]


def extract_docx_text(path: Path) -> ExtractedDocument:
    if docx is None:
        raise RuntimeError("python-docx unavailable")
    document = docx.Document(str(path))
    lines: list[str] = []
    title_hint = path.stem
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name and paragraph.style.name.lower().startswith("title"):
            title_hint = text
        if paragraph.style and paragraph.style.name and paragraph.style.name.lower().startswith("heading"):
            level_match = re.search(r"(\d+)", paragraph.style.name)
            level = int(level_match.group(1)) if level_match else 2
            lines.append("#" * max(1, min(level + 1, 6)) + " " + text)
        else:
            lines.append(text)
    for table in document.tables:
        row_lines = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                row_lines.append(" | ".join(cells))
        if row_lines:
            lines.append("## Table")
            lines.extend(row_lines)
    text = normalize_text("\n\n".join(lines))
    return ExtractedDocument(text=text, title_hint=title_hint)


def extract_pdf_text(path: Path) -> ExtractedDocument:
    if pdfplumber is None:
        raise RuntimeError("pdfplumber unavailable")
    pages: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for idx, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            if not text:
                continue
            pages.append(f"## Page {idx}\n{text}")
    text = normalize_text("\n\n".join(pages))
    return ExtractedDocument(text=text, title_hint=path.stem)


def extract_document(file: CandidateFile) -> ExtractedDocument:
    suffix = file.path.suffix.lower()
    if suffix == ".md":
        return ExtractedDocument(
            text=normalize_text(file.path.read_text(encoding="utf-8", errors="ignore")),
            title_hint=file.path.stem,
        )
    if suffix == ".docx":
        return extract_docx_text(file.path)
    if suffix == ".pdf":
        return extract_pdf_text(file.path)
    raise RuntimeError(f"unsupported suffix: {suffix}")


def derive_priority_layer(path: Path | str) -> str:
    """Derive priority_layer from absolute or relative path string.

    L1: /03_delivery/  L2: /knowledge/  L3: /02_draft/  L4: /00_source/  L5: 其他
    """
    posix = (path.as_posix() if isinstance(path, Path) else str(path))
    if "/03_delivery/" in posix:
        return "L1"
    if "/knowledge/" in posix or "/dyna/knowledge/" in posix:
        return "L2"
    if "/02_draft/" in posix:
        return "L3"
    if "/00_source/" in posix:
        return "L4"
    return "L5"


_EMBEDDER = None


def get_embedder():
    """Lazy-load BAAI/bge-m3 SentenceTransformer; module-level singleton."""
    global _EMBEDDER
    if _EMBEDDER is not None:
        return _EMBEDDER
    from sentence_transformers import SentenceTransformer  # type: ignore
    _EMBEDDER = SentenceTransformer(EMBEDDING_MODEL_NAME)
    return _EMBEDDER


def encode_chunks(texts: list[str]) -> list[bytes]:
    """Encode chunk texts -> list of float32 BLOBs (1024-dim, L2-normalized)."""
    if not texts:
        return []
    import numpy as np  # type: ignore
    model = get_embedder()
    arr = np.asarray(
        model.encode(
            texts,
            normalize_embeddings=True,
            show_progress_bar=False,
            batch_size=EMBEDDING_BATCH_SIZE,
        ),
        dtype=np.float32,
    )
    return [row.tobytes() for row in arr]


def migrate_schema(conn: sqlite3.Connection) -> None:
    """Idempotent ALTER TABLE: add embedding + priority_layer columns if missing."""
    cols = {row[1] for row in conn.execute("PRAGMA table_info(chunks)").fetchall()}
    if "embedding" not in cols:
        conn.execute("ALTER TABLE chunks ADD COLUMN embedding BLOB")
    if "priority_layer" not in cols:
        conn.execute("ALTER TABLE chunks ADD COLUMN priority_layer TEXT NOT NULL DEFAULT 'L5'")
    conn.commit()


def connect(db_path: Path) -> sqlite3.Connection:
    db_path.parent.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    conn.executescript(SCHEMA_SQL)
    migrate_schema(conn)
    return conn


def delete_doc_chunks(conn: sqlite3.Connection, doc_id: int) -> None:
    row_ids = [row[0] for row in conn.execute("SELECT id FROM chunks WHERE doc_id = ?", (doc_id,))]
    if row_ids:
        conn.executemany("DELETE FROM chunks_fts WHERE rowid = ?", [(row_id,) for row_id in row_ids])
    conn.execute("DELETE FROM chunks WHERE doc_id = ?", (doc_id,))


def upsert_document(conn: sqlite3.Connection, file: CandidateFile, text: str, title_hint: str) -> tuple[str, int]:
    stat = file.path.stat()
    content_sha = sha256_text(text)
    title = extract_title(text, title_hint)
    indexed_at = now_iso()

    existing = conn.execute(
        "SELECT id, content_sha256 FROM documents WHERE path = ?",
        (str(file.path),),
    ).fetchone()

    if existing and existing["content_sha256"] == content_sha:
        conn.execute(
            "UPDATE documents SET mtime_ns = ?, size_bytes = ?, indexed_at = ?, title = ? WHERE id = ?",
            (stat.st_mtime_ns, stat.st_size, indexed_at, title, existing["id"]),
        )
        return "unchanged", 0

    conn.execute(
        """
        INSERT INTO documents(path, slug, project_dir, doc_type, title, ext, mtime_ns, size_bytes, content_sha256, indexed_at)
        VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(path) DO UPDATE SET
            slug=excluded.slug,
            project_dir=excluded.project_dir,
            doc_type=excluded.doc_type,
            title=excluded.title,
            ext=excluded.ext,
            mtime_ns=excluded.mtime_ns,
            size_bytes=excluded.size_bytes,
            content_sha256=excluded.content_sha256,
            indexed_at=excluded.indexed_at
        """,
        (
            str(file.path),
            file.slug,
            file.project_dir,
            file.doc_type,
            title,
            file.path.suffix.lower(),
            stat.st_mtime_ns,
            stat.st_size,
            content_sha,
            indexed_at,
        ),
    )
    doc_id = conn.execute("SELECT id FROM documents WHERE path = ?", (str(file.path),)).fetchone()[0]
    delete_doc_chunks(conn, doc_id)

    chunks = chunk_text(text)
    priority_layer = derive_priority_layer(file.path)
    embeddings: list[bytes] = []
    if chunks:
        # fail-fast：encoder 不可用直接 raise，让上游 cron 报错 DM
        # 历史问题：silent degrade 写 NULL embedding 导致 717 chunks 静默失明（2026-05-10 修复）
        embeddings = encode_chunks([content for _section, content in chunks])
    fts_rows = []
    for chunk_index, (section, content) in enumerate(chunks):
        emb = embeddings[chunk_index] if chunk_index < len(embeddings) and embeddings[chunk_index] else None
        cursor = conn.execute(
            """
            INSERT INTO chunks(doc_id, chunk_index, path, slug, doc_type, title, section, content, embedding, priority_layer)
            VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (doc_id, chunk_index, str(file.path), file.slug, file.doc_type, title, section, content, emb, priority_layer),
        )
        row_id = cursor.lastrowid
        fts_rows.append((row_id, str(file.path), file.slug, file.doc_type, title, section, content))

    if fts_rows:
        conn.executemany(
            "INSERT INTO chunks_fts(rowid, path, slug, doc_type, title, section, content) VALUES(?, ?, ?, ?, ?, ?, ?)",
            fts_rows,
        )
    return ("updated" if existing else "inserted"), len(fts_rows)


def delete_stale_documents(conn: sqlite3.Connection, live_paths: set[str], slug_filter: str | None = None) -> int:
    if slug_filter:
        stale = conn.execute("SELECT id, path FROM documents WHERE slug = ?", (slug_filter,)).fetchall()
    else:
        stale = conn.execute("SELECT id, path FROM documents").fetchall()
    deleted = 0
    for row in stale:
        if row["path"] in live_paths:
            continue
        delete_doc_chunks(conn, row["id"])
        conn.execute("DELETE FROM documents WHERE id = ?", (row["id"],))
        deleted += 1
    return deleted


def semantic_rerank(
    conn: sqlite3.Connection,
    query: str,
    slug: str | None,
    doc_type: str | None,
    limit: int,
    pool: int = 50,
    bm25_weight: float = 0.4,
    cosine_weight: float = 0.6,
) -> list[dict]:
    """FTS5 top-`pool` 召回 + bge-m3 cosine rerank，返回 top-`limit`。

    任何阶段失败（无 embedding 列 / 编码失败 / FTS5 异常）返回空列表，由调用方走纯 FTS 回退。
    """
    if len(query.strip()) < 3:
        return []
    try:
        import numpy as np  # type: ignore
    except Exception:
        return []
    filters: list[str] = []
    params: list[object] = [query]
    if slug:
        filters.append("c.slug = ?")
        params.append(slug)
    if doc_type:
        filters.append("c.doc_type = ?")
        params.append(doc_type)
    where = (" AND " + " AND ".join(filters)) if filters else ""
    params.append(pool)
    try:
        rows = conn.execute(
            f"""
            SELECT
                c.id AS chunk_id,
                c.path, c.slug, c.doc_type, d.ext, c.title, c.section,
                c.embedding, c.priority_layer,
                snippet(chunks_fts, 5, '[', ']', ' … ', 18) AS snippet,
                bm25(chunks_fts, 3.5, 2.0, 1.8, 1.8, 1.5, 1.0) AS bm25
            FROM chunks_fts
            JOIN chunks c ON c.id = chunks_fts.rowid
            JOIN documents d ON d.id = c.doc_id
            WHERE chunks_fts MATCH ? {where}
            ORDER BY bm25
            LIMIT ?
            """,
            params,
        ).fetchall()
    except Exception:
        return []
    if not rows:
        return []
    candidates = [dict(r) for r in rows]
    # encode query
    try:
        q_vec = encode_chunks([query])
        if not q_vec or not q_vec[0]:
            return []
        q_arr = np.frombuffer(q_vec[0], dtype=np.float32)
    except Exception:
        return []
    # bm25 normalization (smaller = better in FTS5; flip sign)
    bm25_vals = [float(c["bm25"]) for c in candidates]
    bm25_neg = [-v for v in bm25_vals]
    bmin, bmax = min(bm25_neg), max(bm25_neg)
    span = (bmax - bmin) or 1.0
    out: list[dict] = []
    for c, bn in zip(candidates, bm25_neg, strict=True):
        emb = c.get("embedding")
        if emb:
            try:
                v = np.frombuffer(emb, dtype=np.float32)
                cos = float(np.dot(q_arr, v))  # both L2-normalized
            except Exception:
                cos = 0.0
        else:
            cos = 0.0
        bm_norm = (bn - bmin) / span
        combined = bm25_weight * bm_norm + cosine_weight * cos
        out.append({
            "path": c["path"],
            "slug": c["slug"],
            "doc_type": c["doc_type"],
            "ext": c["ext"],
            "title": c["title"],
            "section": c["section"],
            "snippet": c["snippet"],
            "priority_layer": c["priority_layer"],
            "bm25": bm25_vals[candidates.index(c)] if False else float(c["bm25"]),
            "cosine": cos,
            "score": -combined,  # negative so smaller = better (matches FTS convention)
        })
    out.sort(key=lambda r: r["score"])
    return out[:limit]


def search_rows(conn: sqlite3.Connection, query: str, slug: str | None, doc_type: str | None, limit: int) -> list[dict]:
    # Trigram tokenizer requires >= 3 chars for MATCH; fall back to LIKE for short queries
    if len(query.strip()) < 3:
        return _search_rows_like(conn, query, slug, doc_type, limit)

    filters = []
    params: list[object] = [query]
    if slug:
        filters.append("c.slug = ?")
        params.append(slug)
    if doc_type:
        filters.append("c.doc_type = ?")
        params.append(doc_type)
    where = ""
    if filters:
        where = " AND " + " AND ".join(filters)
    params.append(limit)
    try:
        rows = conn.execute(
            f"""
            SELECT
                c.path,
                c.slug,
                c.doc_type,
                d.ext,
                c.title,
                c.section,
                snippet(chunks_fts, 5, '[', ']', ' … ', 18) AS snippet,
                c.priority_layer AS priority_layer,
                bm25(chunks_fts, 3.5, 2.0, 1.8, 1.8, 1.5, 1.0) * CASE c.doc_type
                    WHEN '03_delivery' THEN 1.5
                    WHEN 'published'   THEN 1.3
                    WHEN '00_source'   THEN 1.2
                    WHEN '01_meetings' THEN 1.0
                    WHEN 'reference'   THEN 1.0
                    WHEN '02_draft'    THEN 0.6
                    ELSE 1.0
                END AS score
            FROM chunks_fts
            JOIN chunks c ON c.id = chunks_fts.rowid
            JOIN documents d ON d.id = c.doc_id
            WHERE chunks_fts MATCH ? {where}
            ORDER BY score
            LIMIT ?
            """,
            params,
        ).fetchall()
    except Exception:
        return []
    results = [dict(row) for row in rows]
    fallback = _search_rows_like_terms(conn, query, slug, doc_type, limit) if len(_query_like_terms(query)) >= 3 else []
    if results and not fallback:
        return results
    if not results:
        return fallback
    merged: dict[tuple[object, object], dict] = {}
    for row in results + fallback:
        key = (row.get("path"), row.get("section"))
        existing = merged.get(key)
        if not existing or row.get("score", 0) < existing.get("score", 0):
            merged[key] = row
    return sorted(merged.values(), key=lambda row: row.get("score", 0))[:limit]


def _search_rows_like(conn: sqlite3.Connection, query: str, slug: str | None, doc_type: str | None, limit: int) -> list[dict]:
    """Fallback for queries too short for trigram FTS5 (< 3 chars)."""
    filters = ["c.content LIKE ?"]
    params: list[object] = [f"%{query}%"]
    if slug:
        filters.append("c.slug = ?")
        params.append(slug)
    if doc_type:
        filters.append("c.doc_type = ?")
        params.append(doc_type)
    where = " AND ".join(filters)
    params.append(limit)
    try:
        rows = conn.execute(
            f"""
            SELECT
                c.path,
                c.slug,
                c.doc_type,
                d.ext,
                c.title,
                c.section,
                c.content AS snippet,
                c.priority_layer AS priority_layer,
                1.0 AS score
            FROM chunks c
            JOIN documents d ON d.id = c.doc_id
            WHERE {where}
            ORDER BY c.doc_type DESC
            LIMIT ?
            """,
            params,
        ).fetchall()
    except Exception:
        return []
    return [dict(row) for row in rows]


def _query_like_terms(query: str) -> list[str]:
    terms: list[str] = []
    seen: set[str] = set()

    def add(term: str) -> None:
        term = term.strip()
        if len(term) < 2:
            return
        key = term.lower()
        if key in seen:
            return
        seen.add(key)
        terms.append(term)

    for token in re.findall(r"[A-Za-z0-9][A-Za-z0-9+._-]*|[\u4e00-\u9fff]+|[ぁ-んァ-ヶー]+", query):
        add(token)
        if re.fullmatch(r"[\u4e00-\u9fff]+", token):
            for width in (2, 3):
                if len(token) <= width:
                    continue
                for idx in range(0, len(token) - width + 1):
                    add(token[idx : idx + width])
    return terms[:32]


def _search_rows_like_terms(conn: sqlite3.Connection, query: str, slug: str | None, doc_type: str | None, limit: int) -> list[dict]:
    terms = _query_like_terms(query)
    if not terms:
        return []

    filters = []
    params: list[object] = []
    if slug:
        filters.append("c.slug = ?")
        params.append(slug)
    if doc_type:
        filters.append("c.doc_type = ?")
        params.append(doc_type)
    match_exprs = []
    match_params: list[object] = []
    score_exprs = []
    score_params: list[object] = []
    for term in terms:
        like = f"%{term}%"
        match_exprs.append("(c.title LIKE ? OR c.section LIKE ? OR c.content LIKE ?)")
        match_params.extend([like, like, like])
        score_exprs.append("(CASE WHEN c.title LIKE ? THEN 3 ELSE 0 END + CASE WHEN c.section LIKE ? THEN 2 ELSE 0 END + CASE WHEN c.content LIKE ? THEN 1 ELSE 0 END)")
        score_params.extend([like, like, like])
    where = " AND ".join(filters + ["(" + " OR ".join(match_exprs) + ")"])
    score_sql = " + ".join(score_exprs)
    sql_params = score_params + params + match_params + [limit]
    try:
        rows = conn.execute(
            f"""
            SELECT
                c.path,
                c.slug,
                c.doc_type,
                d.ext,
                c.title,
                c.section,
                c.content AS snippet,
                c.priority_layer AS priority_layer,
                -1.0 * ({score_sql}) AS score
            FROM chunks c
            JOIN documents d ON d.id = c.doc_id
            WHERE {where}
            ORDER BY score, c.doc_type DESC
            LIMIT ?
            """,
            sql_params,
        ).fetchall()
    except Exception:
        return []
    return [dict(row) for row in rows]


def command_update(args: argparse.Namespace) -> int:
    db_path = Path(args.db_path).expanduser()
    conn = connect(db_path)
    inserted = updated = unchanged = chunk_total = skipped_empty = skipped_extract = 0
    live_paths: set[str] = set()
    slug_filter = args.slug

    try:
        for file in iter_candidate_files(slug_filter=slug_filter):
            live_paths.add(str(file.path))
            try:
                extracted = extract_document(file)
            except Exception:
                skipped_extract += 1
                continue
            if not extracted.text.strip():
                skipped_empty += 1
                continue
            status, chunk_count = upsert_document(conn, file, extracted.text, extracted.title_hint)
            chunk_total += chunk_count
            if status == "inserted":
                inserted += 1
            elif status == "updated":
                updated += 1
            else:
                unchanged += 1

        deleted = delete_stale_documents(conn, live_paths, slug_filter=slug_filter)
        conn.execute(
            "INSERT INTO index_meta(key, value) VALUES('last_updated_at', ?) ON CONFLICT(key) DO UPDATE SET value=excluded.value",
            (now_iso(),),
        )
        if slug_filter:
            conn.execute(
                "INSERT INTO index_meta(key, value) VALUES(?, ?) ON CONFLICT(key) DO UPDATE SET value=excluded.value",
                (f"last_updated_at:{slug_filter}", now_iso()),
            )
        conn.execute(
            "INSERT INTO index_meta(key, value) VALUES('projects_root', ?) ON CONFLICT(key) DO UPDATE SET value=excluded.value",
            (str(DOC_ROOT or PROJECTS_ROOT),),
        )
        conn.commit()

        docs = conn.execute("SELECT COUNT(*) FROM documents").fetchone()[0]
        chunks = conn.execute("SELECT COUNT(*) FROM chunks").fetchone()[0]
        null_emb = conn.execute("SELECT COUNT(*) FROM chunks WHERE embedding IS NULL").fetchone()[0]
        payload = {
            "db_path": str(db_path),
            "slug": slug_filter,
            "documents": docs,
            "chunks": chunks,
            "null_embedding_count": null_emb,
            "inserted": inserted,
            "updated": updated,
            "unchanged": unchanged,
            "deleted": deleted,
            "chunk_writes": chunk_total,
            "skipped_empty": skipped_empty,
            "skipped_extract": skipped_extract,
        }
        if args.json:
            print(json.dumps(payload, ensure_ascii=False, indent=2))
        else:
            print(f"db={db_path}")
            print(f"documents={docs} chunks={chunks} null_embedding_count={null_emb}")
            print(
                "inserted={inserted} updated={updated} unchanged={unchanged} deleted={deleted} skipped_empty={skipped_empty} skipped_extract={skipped_extract}".format(
                    inserted=inserted,
                    updated=updated,
                    unchanged=unchanged,
                    deleted=deleted,
                    skipped_empty=skipped_empty,
                    skipped_extract=skipped_extract,
                )
            )
        # 后置断言：null_embedding_count > 0 视为部分失败，让上游 cron 报错 DM 提醒跑 rebuild
        if null_emb > 0:
            print(f"WARN: {null_emb} chunks have NULL embedding; run workspace-doc-rebuild-embeddings.sh", file=sys.stderr)
            return 2
        return 0
    finally:
        conn.close()


def command_search(args: argparse.Namespace) -> int:
    db_path = Path(args.db_path).expanduser()
    conn = connect(db_path)
    try:
        results: list[dict] = []
        if not getattr(args, "no_semantic", False):
            results = semantic_rerank(conn, args.query, args.slug, args.doc_type, args.limit)
        if not results:
            results = search_rows(conn, args.query, args.slug, args.doc_type, args.limit)
        if args.json:
            print(json.dumps(results, ensure_ascii=False, indent=2))
        else:
            for i, row in enumerate(results, start=1):
                print(f"[{i}] {row['slug']} {row['doc_type']} {row['title']}")
                print(f"    path={row['path']}")
                print(f"    ext={row['ext']} section={row['section']} score={row['score']:.4f}")
                print(f"    {row['snippet']}")
        return 0
    finally:
        conn.close()


def command_stats(args: argparse.Namespace) -> int:
    db_path = Path(args.db_path).expanduser()
    conn = connect(db_path)
    try:
        docs = conn.execute("SELECT COUNT(*) FROM documents").fetchone()[0]
        chunks = conn.execute("SELECT COUNT(*) FROM chunks").fetchone()[0]
        last_updated = conn.execute("SELECT value FROM index_meta WHERE key = 'last_updated_at'").fetchone()
        by_slug = conn.execute(
            "SELECT slug, COUNT(*) AS docs FROM documents GROUP BY slug ORDER BY slug"
        ).fetchall()
        by_type = conn.execute(
            "SELECT doc_type, COUNT(*) AS docs FROM documents GROUP BY doc_type ORDER BY doc_type"
        ).fetchall()
        by_ext = conn.execute(
            "SELECT ext, COUNT(*) AS docs FROM documents GROUP BY ext ORDER BY ext"
        ).fetchall()
        payload = {
            "db_path": str(db_path),
            "documents": docs,
            "chunks": chunks,
            "last_updated_at": last_updated[0] if last_updated else None,
            "by_slug": [dict(row) for row in by_slug],
            "by_doc_type": [dict(row) for row in by_type],
            "by_ext": [dict(row) for row in by_ext],
        }
        if args.json:
            print(json.dumps(payload, ensure_ascii=False, indent=2))
        else:
            print(f"db={db_path}")
            print(f"documents={docs} chunks={chunks}")
            print(f"last_updated_at={payload['last_updated_at']}")
            print("by_slug:")
            for row in payload["by_slug"]:
                print(f"  {row['slug']}: {row['docs']}")
            print("by_doc_type:")
            for row in payload["by_doc_type"]:
                print(f"  {row['doc_type']}: {row['docs']}")
            print("by_ext:")
            for row in payload["by_ext"]:
                print(f"  {row['ext']}: {row['docs']}")
        return 0
    finally:
        conn.close()


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Workspace document index for SQLite FTS5.")
    parser.add_argument("--db-path", default=str(DEFAULT_DB_PATH or ""), help="SQLite DB path (required via DOC_INDEX_DB env or --db-path)")
    sub = parser.add_subparsers(dest="command", required=True)

    p_update = sub.add_parser("update", help="Incrementally update doc index")
    p_update.add_argument("--slug")
    p_update.add_argument("--json", action="store_true")
    p_update.set_defaults(func=command_update)

    p_search = sub.add_parser("search", help="Search indexed chunks")
    p_search.add_argument("query")
    p_search.add_argument("--slug")
    p_search.add_argument("--doc-type", choices=sorted(INCLUDE_DIR_EXTS.keys()))
    p_search.add_argument("--limit", type=int, default=8)
    p_search.add_argument("--json", action="store_true")
    p_search.add_argument("--no-semantic", action="store_true", help="Disable bge-m3 cosine rerank (FTS5 only)")
    p_search.set_defaults(func=command_search)

    p_stats = sub.add_parser("stats", help="Show index stats")
    p_stats.add_argument("--json", action="store_true")
    p_stats.set_defaults(func=command_stats)
    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    return args.func(args)


if __name__ == "__main__":
    raise SystemExit(main())
